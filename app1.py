# pip install streamlit azure-storage-blob langchain_community langchain-google-genai docx PyPDF2 faiss-cpu numpy requests

import streamlit as st
from azure.storage.blob import BlobServiceClient
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.docstore.document import Document
import faiss
import numpy as np
import tempfile, os, re, pickle
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ==========================
# Streamlit UI
# ==========================
st.title("📄 Azure PDF/DOCX/TXT Q&A with Google Generative AI")

# ==========================
# Config
# ==========================
AZURE_CONNECTION_STRING = st.secrets.get("AZURE_CONNECTION_STRING")  # Store in Streamlit secrets
CONTAINER_NAME = "kalyanpdf"
API_KEY = st.secrets.get("GOOGLE_API_KEY")  # Store in Streamlit secrets

container_client = BlobServiceClient.from_connection_string(
    AZURE_CONNECTION_STRING
).get_container_client(CONTAINER_NAME)

# ==========================
# Text cleaning function
# ==========================
def clean_text(text):
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# ==========================
# Load documents from Azure Blob
# ==========================
@st.cache_data(show_spinner=True)
def load_documents():
    documents_text = []
    for blob in container_client.list_blobs():
        if not blob.name.lower().endswith((".pdf", ".docx", ".txt")):
            continue

        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(blob.name)[1])
        tmp_file.write(container_client.get_blob_client(blob.name).download_blob().readall())
        tmp_file.close()

        ext = os.path.splitext(blob.name)[1].lower()
        if ext == ".pdf":
            text = "\n".join([p.page_content for p in PyPDFLoader(tmp_file.name).load()])
        elif ext == ".docx":
            doc = DocxDocument(tmp_file.name)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".txt":
            with open(tmp_file.name, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        os.remove(tmp_file.name)
        documents_text.append(clean_text(text))
    return documents_text

documents_text = load_documents()
st.write(f"Loaded {len(documents_text)} documents from Azure Blob")

# ==========================
# Split into chunks
# ==========================
splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
all_chunks = []
for txt in documents_text:
    all_chunks.extend(splitter.split_text(txt))
st.write(f"Total chunks created: {len(all_chunks)}")

# ==========================
# Embeddings + FAISS
# ==========================
@st.cache_resource(show_spinner=True)
def build_faiss_index(chunks):
    embedding_model = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=API_KEY
    )
    documents = [Document(page_content=c) for c in chunks]
    embeddings = embedding_model.embed_documents([d.page_content for d in documents])
    embeddings_np = np.array(embeddings, dtype="float32")
    faiss_index = faiss.IndexFlatL2(embeddings_np.shape[1])
    faiss_index.add(embeddings_np)
    return faiss_index, chunks, embedding_model

faiss_index, chunks, embedding_model = build_faiss_index(all_chunks)

# ==========================
# Helpers
# ==========================
def embed_query(q):
    return np.array(embedding_model.embed_query(q), dtype="float32")

def retrieve(q_emb, index, chunks, k=1):
    _, I = index.search(q_emb.reshape(1, -1), k)
    return [chunks[i] for i in I[0]]

def build_prompt(context, question):
    return f"""
Use ONLY the context to answer.

Context:
{context}

Question: {question}
Answer:
"""

def safe_post(url, headers, json_data):
    session = requests.Session()
    retry = Retry(total=3, backoff_factor=0.5)
    session.mount("https://", HTTPAdapter(max_retries=retry))
    return session.post(url, headers=headers, json=json_data, timeout=30)

def generate_answer(context, question):
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
    data = {
        "contents": [{"parts": [{"text": build_prompt(context, question)}]}]
    }
    res = safe_post(f"{url}?key={API_KEY}", {"Content-Type": "application/json"}, data)
    return res.json()["candidates"][0]["content"]["parts"][0]["text"]

# ==========================
# Q&A Interface
# ==========================
query = st.text_input("Ask a question about the documents:")

if st.button("Generate Answer") and query:
    with st.spinner("Generating answer..."):
        q_emb = embed_query(query)
        top_chunk = retrieve(q_emb, faiss_index, chunks, k=1)
        context = "\n\n".join(top_chunk)
        answer = generate_answer(context, query)
        st.markdown(f"**Answer:** {answer}")
