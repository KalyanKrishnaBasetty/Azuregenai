from azure.storage.blob import BlobServiceClient
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.docstore.document import Document
import faiss
import numpy as np
import tempfile, os, re, pickle
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ==========================
# Azure Blob Config
# ==========================
AZURE_CONNECTION_STRING = "DefaultEndpointsProtocol=https;AccountName=azurestoragepdf;AccountKey=JwJYpcDOxsZVA+PmeuMfx838rHqh/gwKg1gsi5b1qGHB1sakfhGuG8fo9WECGrKUGfhw8StGpTlh+ASt1kDjcA==;EndpointSuffix=core.windows.net"
CONTAINER_NAME = "kalyanpdf"
container_client = BlobServiceClient.from_connection_string(
    AZURE_CONNECTION_STRING
).get_container_client(CONTAINER_NAME)

# ==========================
# Text cleaning
# ==========================
def clean_text(text):
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# ==========================
# Load documents (PDF + DOCX + TXT)
# ==========================
documents_text = []

for blob in container_client.list_blobs():
    if not blob.name.lower().endswith((".pdf", ".docx", ".txt")):
        continue

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(blob.name)[1])
    tmp_file.write(container_client.get_blob_client(blob.name).download_blob().readall())
    tmp_file.close()

    ext = os.path.splitext(blob.name)[1].lower()

    if ext == ".pdf":
        text = "\n".join([p.page_content for p in PyPDFLoader(tmp_file.name).load()])

    elif ext == ".docx":
        doc = DocxDocument(tmp_file.name)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif ext == ".txt":
        with open(tmp_file.name, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    os.remove(tmp_file.name)
    documents_text.append(clean_text(text))

# ==========================
# Chunking
# ==========================
splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
all_chunks = []
for txt in documents_text:
    all_chunks.extend(splitter.split_text(txt))

# ==========================
# Embeddings + FAISS
# ==========================
API_KEY = "AIzaSyAbVjU3k7xMPAw09qr4tQZ84JZAfoRp2XM"

embedding_model = GoogleGenerativeAIEmbeddings(
    model="models/embedding-001",
    google_api_key=API_KEY
)

documents = [Document(page_content=c) for c in all_chunks]
embeddings = embedding_model.embed_documents([d.page_content for d in documents])
embeddings_np = np.array(embeddings, dtype="float32")

faiss_index = faiss.IndexFlatL2(embeddings_np.shape[1])
faiss_index.add(embeddings_np)

faiss.write_index(faiss_index, "faiss_index.idx")
with open("chunks.pkl", "wb") as f:
    pickle.dump(all_chunks, f)

# ==========================
# Helpers
# ==========================
def load_faiss():
    return faiss.read_index("faiss_index.idx"), pickle.load(open("chunks.pkl", "rb"))

def embed_query(q):
    return np.array(
        embedding_model.embed_query(q),
        dtype="float32"
    )

def retrieve(q_emb, index, chunks, k=1):
    _, I = index.search(q_emb.reshape(1, -1), k)
    return [chunks[i] for i in I[0]]

def build_prompt(context, question):
    return f"""
Use ONLY the context to answer.

Context:
{context}

Question: {question}
Answer:
"""

def safe_post(url, headers, json_data):
    session = requests.Session()
    retry = Retry(total=3, backoff_factor=0.5)
    session.mount("https://", HTTPAdapter(max_retries=retry))
    return session.post(url, headers=headers, json=json_data, timeout=30)

def generate_answer(context, question):
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
    data = {
        "contents": [{"parts": [{"text": build_prompt(context, question)}]}]
    }
    res = safe_post(f"{url}?key={API_KEY}", {"Content-Type": "application/json"}, data)
    return res.json()["candidates"][0]["content"]["parts"][0]["text"]

# ==========================
# QUESTIONS & ANSWERS
# ==========================
faiss_index, chunks = load_faiss()

questions = [
    "What is React JS?",
    "What is supervised learning?",
    "What is Python?"
]

for q in questions:
    q_emb = embed_query(q)
    top_chunk = retrieve(q_emb, faiss_index, chunks, k=1)
    context = "\n\n".join(top_chunk)
    ans = generate_answer(context, q)

    print("\nQ:", q)
    print("A:", ans)
